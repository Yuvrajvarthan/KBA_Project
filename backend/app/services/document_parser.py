import pdfplumber
import docx
from typing import Optional, Dict, Any
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentParser:
    """Service for parsing different document types"""
    
    @staticmethod
    async def parse_pdf(file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        try:
            text_content = []
            metadata = {}
            
            with pdfplumber.open(file_path) as pdf:
                metadata = {
                    "title": pdf.metadata.get("Title", ""),
                    "author": pdf.metadata.get("Author", ""),
                    "creator": pdf.metadata.get("Creator", ""),
                    "producer": pdf.metadata.get("Producer", ""),
                    "creation_date": str(pdf.metadata.get("CreationDate", "")),
                    "pages": len(pdf.pages)
                }
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append({
                                "page": page_num,
                                "text": page_text.strip()
                            })
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {e}")
                        continue
            
            full_text = "\n\n".join([page["text"] for page in text_content])
            
            return {
                "success": True,
                "text": full_text,
                "metadata": metadata,
                "pages": text_content
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "metadata": {},
                "pages": []
            }
    
    @staticmethod
    async def parse_docx(file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                "pages": len(doc.paragraphs)
            }
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            full_text = "\n\n".join(paragraphs)
            
            return {
                "success": True,
                "text": full_text,
                "metadata": metadata,
                "paragraphs": paragraphs
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "metadata": {},
                "paragraphs": []
            }
    
    @staticmethod
    async def parse_txt(file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Get basic file metadata
            file_path_obj = Path(file_path)
            metadata = {
                "title": file_path_obj.stem,
                "size": file_path_obj.stat().st_size,
                "created": str(file_path_obj.stat().st_ctime),
                "modified": str(file_path_obj.stat().st_mtime)
            }
            
            return {
                "success": True,
                "text": content,
                "metadata": metadata
            }
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                
                file_path_obj = Path(file_path)
                metadata = {
                    "title": file_path_obj.stem,
                    "size": file_path_obj.stat().st_size,
                    "created": str(file_path_obj.stat().st_ctime),
                    "modified": str(file_path_obj.stat().st_mtime)
                }
                
                return {
                    "success": True,
                    "text": content,
                    "metadata": metadata
                }
            except Exception as e:
                logger.error(f"Error parsing TXT {file_path} with latin-1: {e}")
                return {
                    "success": False,
                    "error": str(e),
                    "text": "",
                    "metadata": {}
                }
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "metadata": {}
            }
